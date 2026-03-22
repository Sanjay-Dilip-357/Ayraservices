import os
import re
import csv
import json
import random
import zipfile
from io import BytesIO
from flask import abort
from pathlib import Path
from docx import Document
from functools import wraps
from datetime import datetime
from otp_service import OTPService
from config import (TEMPLATE_CONFIG, RELATION_MAPPING, CAST_OPTIONS, 
                   DEFAULT_USER_PASSWORD, DEFAULT_ADMIN_PASSWORD,
                   SQLALCHEMY_DATABASE_URI, IS_PRODUCTION)
from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file
from models import db, User, Draft, PhoneTracking, init_db

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'ayra_services_secret_key_2026')

# Database configuration - use URI from config
app.config['SQLALCHEMY_DATABASE_URI'] = SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,
    'pool_recycle': 300,
    'pool_size': 10,
    'max_overflow': 20
}


# Initialize database
init_db(app)

PHONE_CSV_FILE = 'phone_numbers.csv'

# Create all template folders if not exist
for template_key, template_info in TEMPLATE_CONFIG.items():
    Path(template_info['folder']).mkdir(parents=True, exist_ok=True)
    if 'unmarried_subfolder' in template_info:
        Path(template_info['unmarried_subfolder']).mkdir(parents=True, exist_ok=True)

GENDER_PRONOUNS = {
    'son': {'HE_SHE': 'he'},
    'daughter': {'HE_SHE': 'she'}
}

GENDER_PRONOUNS_BY_GENDER = {
    'male': {'HE_SHE': 'he'},
    'female': {'HE_SHE': 'she'}
}

# ==================== AUTHENTICATION DECORATORS ====================

def login_required(f):
    """Decorator to require login"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            if request.is_json:
                return jsonify({'success': False, 'message': 'Login required', 'redirect': '/?show_login=true'})
            return redirect(url_for('index', show_login='true'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    """Decorator to require admin or super_admin login"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            if request.is_json:
                return jsonify({'success': False, 'message': 'Login required', 'redirect': '/login'})
            return redirect(url_for('login'))
        if session.get('user_role') not in ['admin', 'super_admin']:
            if request.is_json:
                return jsonify({'success': False, 'message': 'Admin access required'})
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

def super_admin_required(f):
    """Decorator to require super_admin login"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            if request.is_json:
                return jsonify({'success': False, 'message': 'Login required', 'redirect': '/login'})
            return redirect(url_for('login'))
        if session.get('user_role') != 'super_admin':
            if request.is_json:
                return jsonify({'success': False, 'message': 'Super Admin access required'})
            return redirect(url_for('admin_dashboard'))
        return f(*args, **kwargs)
    return decorated_function

# ==================== HELPER: GET DASHBOARD BY ROLE ====================

def get_dashboard_redirect():
    """Get appropriate dashboard redirect based on user role"""
    if 'user_id' not in session:
        return redirect(url_for('index'))
    
    user_role = session.get('user_role')
    
    if user_role == 'super_admin':
        return redirect(url_for('super_admin_dashboard'))
    elif user_role == 'admin':
        return redirect(url_for('admin_dashboard'))
    else:
        return redirect(url_for('dashboard'))

@app.route('/setup-super-admin', methods=['GET', 'POST'])
def setup_super_admin():
    """One-time super admin setup"""

    # 🚨 Block if already exists
    if User.super_admin_exists():
        abort(403)

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip().lower()
        phone = request.form.get('phone', '').strip()
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()

        errors = []

        if not name:
            errors.append("Name is required")
        if not email:
            errors.append("Email is required")
        if not password:
            errors.append("Password is required")
        if len(password) < 8:
            errors.append("Password must be at least 8 characters")
        if password != confirm_password:
            errors.append("Passwords do not match")
        if User.get_by_email(email):
            errors.append("Email already exists")

        if errors:
            return render_template('setup_super_admin.html',
                                   errors=errors,
                                   name=name,
                                   email=email,
                                   phone=phone)

        try:
            super_admin = User(
                name=name,
                email=email,
                phone=phone,
                role='super_admin',
                is_active=True
            )
            super_admin.set_password(password)

            db.session.add(super_admin)
            db.session.commit()

            return render_template('setup_success.html',
                                   title="Super Admin Created!",
                                   name=name,
                                   email=email,
                                   role="Super Administrator",
                                   redirect_url="/")

        except Exception as e:
            db.session.rollback()
            return render_template('setup_super_admin.html',
                                   errors=[str(e)],
                                   name=name,
                                   email=email,
                                   phone=phone)

    return render_template('setup_super_admin.html')

@app.route('/create-super-admin', methods=['GET', 'POST'])
@super_admin_required
def create_super_admin():
    """Create additional super admins (only by super admin)"""

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip().lower()
        phone = request.form.get('phone', '').strip()
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()

        errors = []

        if not name:
            errors.append("Name is required")
        if not email:
            errors.append("Email is required")
        if not password:
            errors.append("Password is required")
        if len(password) < 8:
            errors.append("Password must be at least 8 characters")
        if password != confirm_password:
            errors.append("Passwords do not match")
        if User.get_by_email(email):
            errors.append("Email already exists")

        if errors:
            return render_template('setup_super_admin.html', errors=errors)

        try:
            new_super_admin = User(
                name=name,
                email=email,
                phone=phone,
                role='super_admin',
                is_active=True,
                created_by=session.get('user_id')  # optional tracking
            )
            new_super_admin.set_password(password)

            db.session.add(new_super_admin)
            db.session.commit()

            return render_template('setup_success.html',
                                   title="Super Admin Created!",
                                   name=name,
                                   email=email,
                                   role="Super Administrator",
                                   redirect_url="/superadmin/dashboard")

        except Exception as e:
            db.session.rollback()
            return render_template('setup_super_admin.html',
                                   errors=[str(e)])

    return render_template('setup_super_admin.html')


# ==================== ADMIN SETUP ROUTE (ONE-TIME USE) ====================
@app.route('/setup-admin', methods=['GET', 'POST'])
def setup_admin():
    """Admin setup - requires super admin OR works if no admin exists"""
    
    # Check if user is logged in as super admin
    is_super_admin = session.get('user_role') == 'super_admin'
    
    # If not super admin and an admin already exists, deny access
    if not is_super_admin and User.admin_exists():
        return render_template('setup_complete.html',
                             message="Admin setup requires Super Admin access",
                             redirect_url="/")
    
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip().lower()
        phone = request.form.get('phone', '').strip()
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        
        # Validation
        errors = []
        
        if not name:
            errors.append("Name is required")
        if not email:
            errors.append("Email is required")
        if not password:
            errors.append("Password is required")
        if len(password) < 8:
            errors.append("Password must be at least 8 characters")
        if password != confirm_password:
            errors.append("Passwords do not match")
        if User.get_by_email(email):
            errors.append("Email already exists")
        
        if errors:
            return render_template('setup_admin.html', errors=errors,
                                 name=name, email=email, phone=phone,
                                 is_super_admin=is_super_admin)
        
        try:
            # Create admin user
            admin = User(
                name=name,
                email=email,
                phone=phone,
                role='admin',
                is_active=True,
                created_by=session.get('user_id') if is_super_admin else None
            )
            admin.set_password(password)
            
            db.session.add(admin)
            db.session.commit()
            
            return render_template('setup_success.html',
                                 title="Admin Created!",
                                 name=name,
                                 email=email,
                                 role="Administrator",
                                 redirect_url="/superadmin/dashboard" if is_super_admin else "/")
            
        except Exception as e:
            db.session.rollback()
            return render_template('setup_admin.html',
                                 errors=[f"Error: {str(e)}"],
                                 name=name, email=email, phone=phone,
                                 is_super_admin=is_super_admin)
    
    # GET request - show form
    return render_template('setup_admin.html', is_super_admin=is_super_admin)

@app.route('/superadmin/dashboard')
@super_admin_required
def super_admin_dashboard():
    """Super Admin dashboard"""
    user_id = session.get('user_id')
    super_admin = User.query.get(user_id)
    
    return render_template('super_admin_dashboard.html',
                          admin=super_admin,
                          admin_name=super_admin.name)

@app.route('/api/superadmin/stats')
@super_admin_required
def api_super_admin_stats():
    """Get super admin dashboard statistics"""
    try:
        # Count by role
        super_admin_count = User.query.filter_by(role='super_admin').count()
        admin_count = User.query.filter_by(role='admin').count()
        user_count = User.query.filter_by(role='user').count()
        
        # Active counts
        active_admins = User.query.filter_by(role='admin', is_active=True).count()
        active_users = User.query.filter_by(role='user', is_active=True).count()
        
        # Document stats
        total_docs = Draft.query.count()
        draft_count = Draft.query.filter_by(status='draft').count()
        pending_count = Draft.query.filter_by(status='pending').count()
        approved_count = Draft.query.filter_by(status='approved').count()
        generated_count = Draft.query.filter_by(status='generated').count()
        
        # Get all admins with their stats
        admins = User.query.filter_by(role='admin').all()
        admin_stats = []
        for admin in admins:
            admin_stats.append({
                'id': admin.id,
                'name': admin.name,
                'email': admin.email,
                'phone': admin.phone,
                'is_active': admin.is_active,
                'last_login': admin.last_login.isoformat() if admin.last_login else None,
                'created_at': admin.created_at.isoformat() if admin.created_at else None
            })
        
        # Get all users with their stats
        users = User.query.filter_by(role='user').all()
        user_stats = []
        for user in users:
            user_stats.append({
                'id': user.id,
                'name': user.name,
                'email': user.email,
                'phone': user.phone,
                'is_active': user.is_active,
                'last_login': user.last_login.isoformat() if user.last_login else None,
                'created_at': user.created_at.isoformat() if user.created_at else None,
                'stats': {
                    'drafts': Draft.query.filter_by(user_id=user.id, status='draft').count(),
                    'pending': Draft.query.filter_by(user_id=user.id, status='pending').count(),
                    'approved': Draft.query.filter_by(user_id=user.id, status='approved').count(),
                    'generated': Draft.query.filter_by(user_id=user.id, status='generated').count(),
                    'total': Draft.query.filter_by(user_id=user.id).count()
                }
            })
        
        return jsonify({
            'success': True,
            'overall': {
                'super_admins': super_admin_count,
                'total_admins': admin_count,
                'active_admins': active_admins,
                'total_users': user_count,
                'active_users': active_users,
                'total_documents': total_docs,
                'drafts': draft_count,
                'pending': pending_count,
                'approved': approved_count,
                'generated': generated_count
            },
            'admins': admin_stats,
            'users': user_stats
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})
    
@app.route('/api/superadmin/admins', methods=['GET'])
@super_admin_required
def api_get_admins():
    """Get all admins"""
    try:
        admins = User.query.filter_by(role='admin').order_by(User.created_at.desc()).all()
        return jsonify({
            'success': True,
            'admins': [a.to_dict() for a in admins]
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})
    
@app.route('/api/superadmin/admins', methods=['POST'])
@super_admin_required
def api_create_admin():
    """Create a new admin"""
    try:
        name = request.json.get('name', '').strip()
        email = request.json.get('email', '').strip().lower()
        phone = request.json.get('phone', '').strip()
        password = request.json.get('password', DEFAULT_ADMIN_PASSWORD)
        
        if not name or not email:
            return jsonify({'success': False, 'message': 'Name and email are required'})
        
        if User.get_by_email(email):
            return jsonify({'success': False, 'message': 'Email already exists'})
        
        # Create admin
        admin = User(
            name=name,
            email=email,
            phone=phone,
            role='admin',
            is_active=True,
            created_by=session.get('user_id')
        )
        admin.set_password(password)
        
        db.session.add(admin)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Admin created successfully',
            'admin': admin.to_dict()
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})
    
@app.route('/api/superadmin/admins/<admin_id>', methods=['PUT'])
@super_admin_required
def api_update_admin(admin_id):
    """Update an admin"""
    try:
        admin = User.query.get(admin_id)
        if not admin or admin.role != 'admin':
            return jsonify({'success': False, 'message': 'Admin not found'})
        
        name = request.json.get('name', '').strip()
        phone = request.json.get('phone', '').strip()
        new_email = request.json.get('email', '').strip().lower()
        new_password = request.json.get('password', '').strip()
        
        if name:
            admin.name = name
        if phone is not None:
            admin.phone = phone
        
        if new_email and new_email != admin.email:
            existing = User.get_by_email(new_email)
            if existing and existing.id != admin_id:
                return jsonify({'success': False, 'message': 'Email already exists'})
            admin.email = new_email
        
        if new_password:
            admin.set_password(new_password)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Admin updated successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@app.route('/api/superadmin/admins/<admin_id>/toggle', methods=['POST'])
@super_admin_required
def api_toggle_admin(admin_id):
    """Toggle admin active status"""
    try:
        admin = User.query.get(admin_id)
        if not admin or admin.role != 'admin':
            return jsonify({'success': False, 'message': 'Admin not found'})
        
        admin.is_active = not admin.is_active
        db.session.commit()
        
        status = 'activated' if admin.is_active else 'deactivated'
        return jsonify({'success': True, 'message': f'Admin {status} successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@app.route('/api/superadmin/admins/<admin_id>', methods=['DELETE'])
@super_admin_required
def api_delete_admin(admin_id):
    """Delete an admin"""
    try:
        admin = User.query.get(admin_id)
        if not admin or admin.role != 'admin':
            return jsonify({'success': False, 'message': 'Admin not found'})
        
        db.session.delete(admin)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Admin deleted successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@app.route('/api/superadmin/users', methods=['GET'])
@super_admin_required
def api_super_admin_get_users():
    """Get all users (for super admin)"""
    try:
        users = User.query.filter_by(role='user').order_by(User.created_at.desc()).all()
        return jsonify({
            'success': True,
            'users': [u.to_dict() for u in users]
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


@app.route('/api/superadmin/users', methods=['POST'])
@super_admin_required
def api_super_admin_create_user():
    """Create a new user (by super admin)"""
    try:
        name = request.json.get('name', '').strip()
        email = request.json.get('email', '').strip().lower()
        phone = request.json.get('phone', '').strip()
        password = request.json.get('password', DEFAULT_USER_PASSWORD)
        
        if not name or not email:
            return jsonify({'success': False, 'message': 'Name and email are required'})
        
        if User.get_by_email(email):
            return jsonify({'success': False, 'message': 'Email already exists'})
        
        # Create user
        user = User(
            name=name,
            email=email,
            phone=phone,
            role='user',
            is_active=True,
            created_by=session.get('user_id')
        )
        user.set_password(password)
        
        db.session.add(user)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'User created successfully',
            'user': user.to_dict()
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@app.route('/api/superadmin/users/<user_id>', methods=['PUT'])
@super_admin_required
def api_super_admin_update_user(user_id):
    """Update a user (by super admin)"""
    try:
        user = User.query.get(user_id)
        if not user or user.role not in ['user', 'admin']:
            return jsonify({'success': False, 'message': 'User not found'})
        
        name = request.json.get('name', '').strip()
        phone = request.json.get('phone', '').strip()
        new_email = request.json.get('email', '').strip().lower()
        new_password = request.json.get('password', '').strip()
        
        if name:
            user.name = name
        if phone is not None:
            user.phone = phone
        
        if new_email and new_email != user.email:
            existing = User.get_by_email(new_email)
            if existing and existing.id != user_id:
                return jsonify({'success': False, 'message': 'Email already exists'})
            user.email = new_email
        
        if new_password:
            user.set_password(new_password)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'User updated successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@app.route('/api/superadmin/users/<user_id>/toggle', methods=['POST'])
@super_admin_required
def api_super_admin_toggle_user(user_id):
    """Toggle user active status (by super admin)"""
    try:
        user = User.query.get(user_id)
        if not user or user.role not in ['user', 'admin']:
            return jsonify({'success': False, 'message': 'User not found'})
        
        user.is_active = not user.is_active
        db.session.commit()
        
        status = 'activated' if user.is_active else 'deactivated'
        return jsonify({'success': True, 'message': f'User {status} successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@app.route('/api/superadmin/users/<user_id>', methods=['DELETE'])
@super_admin_required
def api_super_admin_delete_user(user_id):
    """Delete a user (by super admin)"""
    try:
        user = User.query.get(user_id)
        if not user or user.role not in ['user', 'admin']:
            return jsonify({'success': False, 'message': 'User not found'})
        
        db.session.delete(user)
        db.session.commit()
        return jsonify({'success': True, 'message': 'User deleted successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})


@app.route('/api/superadmin/documents')
@super_admin_required
def api_super_admin_get_documents():
    """Get all documents (for super admin)"""
    try:
        status = request.args.get('status')
        
        query = Draft.query.order_by(Draft.modified_at.desc())
        if status:
            query = query.filter_by(status=status)
        
        drafts = query.all()
        return jsonify({
            'success': True,
            'documents': [d.to_dict() for d in drafts]
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

def render_setup_form(error=None):
    """Render the admin setup form"""
    error_html = f'<div class="alert alert-danger">{error}</div>' if error else ''
    
    return f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Admin Setup - AYRA Services</title>
            <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
            <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">
            <style>
                body {{
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    min-height: 100vh;
                    display: flex;
                    align-items: center;
                }}
                .setup-card {{
                    background: white;
                    border-radius: 15px;
                    box-shadow: 0 10px 40px rgba(0,0,0,0.2);
                    padding: 2rem;
                }}
                .setup-header {{
                    text-align: center;
                    margin-bottom: 2rem;
                }}
                .setup-icon {{
                    width: 80px;
                    height: 80px;
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    border-radius: 50%;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    margin: 0 auto 1rem;
                    color: white;
                    font-size: 2rem;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="row justify-content-center">
                    <div class="col-md-6">
                        <div class="setup-card">
                            <div class="setup-header">
                                <div class="setup-icon">
                                    <i class="bi bi-shield-lock-fill"></i>
                                </div>
                                <h2>Admin Setup</h2>
                                <p class="text-muted">Create your administrator account</p>
                            </div>
                            
                            {error_html}
                            
                            <form method="POST">
                                <div class="mb-3">
                                    <label class="form-label">
                                        <i class="bi bi-person me-1"></i>Full Name
                                    </label>
                                    <input type="text" class="form-control" name="name" required
                                           placeholder="Enter your full name">
                                </div>
                                
                                <div class="mb-3">
                                    <label class="form-label">
                                        <i class="bi bi-envelope me-1"></i>Email Address
                                    </label>
                                    <input type="email" class="form-control" name="email" required
                                           placeholder="admin@example.com">
                                </div>
                                
                                <div class="mb-3">
                                    <label class="form-label">
                                        <i class="bi bi-lock me-1"></i>Password
                                    </label>
                                    <input type="password" class="form-control" name="password" required
                                           placeholder="Minimum 8 characters">
                                    <small class="text-muted">Must be at least 8 characters long</small>
                                </div>
                                
                                <div class="mb-4">
                                    <label class="form-label">
                                        <i class="bi bi-lock-fill me-1"></i>Confirm Password
                                    </label>
                                    <input type="password" class="form-control" name="confirm_password" required
                                           placeholder="Re-enter password">
                                </div>
                                
                                <button type="submit" class="btn btn-primary w-100 py-2">
                                    <i class="bi bi-check-circle me-2"></i>Create Admin Account
                                </button>
                            </form>
                            
                            <div class="text-center mt-3">
                                <small class="text-muted">
                                    <i class="bi bi-info-circle me-1"></i>
                                    This page will be disabled after admin creation
                                </small>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </body>
        </html>
    """

@app.route('/login')
def login():
    """Login page - redirect to index since we use modal login now"""
    if 'user_id' in session:
        user_role = session.get('user_role')
        
        # Redirect based on role
        if user_role == 'super_admin':
            return redirect(url_for('super_admin_dashboard'))
        elif user_role == 'admin':
            return redirect(url_for('admin_dashboard'))
        else:
            return redirect(url_for('dashboard'))
    
    return redirect(url_for('index'))



@app.route('/api/auth/login', methods=['POST'])
def api_login():
    """API endpoint for login (handles super_admin, admin, and user)"""
    try:
        email = request.json.get('email', '').strip().lower()
        password = request.json.get('password', '').strip()
        
        if not email or not password:
            return jsonify({'success': False, 'message': 'Email and password are required'}), 400
        
        user = User.get_by_email(email)
        
        if not user:
            return jsonify({'success': False, 'message': 'Invalid email or password'}), 401
        
        if not user.check_password(password):
            return jsonify({'success': False, 'message': 'Invalid email or password'}), 401
        
        if not user.is_active:
            return jsonify({'success': False, 'message': 'Your account has been deactivated. Please contact admin.'}), 403
        
        # Set session
        session['user_id'] = user.id
        session['user_email'] = user.email
        session['user_name'] = user.name
        session['user_role'] = user.role
        
        # Update last login
        user.last_login = datetime.utcnow()
        db.session.commit()
        
        app.logger.info(f'User logged in: {email} (Role: {user.role})')
        
        return jsonify({
            'success': True,
            'message': 'Login successful',
            'user': {
                'id': user.id,
                'name': user.name,
                'email': user.email,
                'role': user.role
            }
        }), 200
    
    except Exception as e:
        app.logger.error(f'Login error: {str(e)}')
        return jsonify({'success': False, 'message': 'An error occurred during login'}), 500

@app.route('/api/auth/logout', methods=['POST'])
def api_logout():
    """API endpoint for logout"""
    session.clear()
    return jsonify({'success': True, 'message': 'Logged out successfully', 'redirect': '/'})

@app.route('/api/auth/check')
def api_check_auth():
    """Check if user is authenticated"""
    if 'user_id' in session:
        return jsonify({
            'success': True,
            'authenticated': True,
            'user': {
                'id': session.get('user_id'),
                'name': session.get('user_name'),
                'email': session.get('user_email'),
                'role': session.get('user_role')
            }
        })
    return jsonify({'success': True, 'authenticated': False})

@app.route('/logout')
def logout():
    """Logout and redirect to index"""
    session.clear()
    return redirect(url_for('index'))

# ==================== USER PROFILE MANAGEMENT ====================
@app.route('/api/user/profile', methods=['GET'])
@login_required
def api_get_user_profile():
    """Get current user profile"""
    try:
        user_id = session.get('user_id')
        user = User.query.get(user_id)
        
        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404
        
        return jsonify({
            'success': True,
            'user': {
                'id': user.id,
                'name': user.name,
                'email': user.email,
                'phone': user.phone or '',
                'role': user.role
            }
        })
    except Exception as e:
        app.logger.error(f'Error getting profile: {str(e)}')
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/user/profile', methods=['PUT'])
@login_required
def api_update_user_profile():
    """Update current user profile"""
    try:
        user_id = session.get('user_id')
        user = User.query.get(user_id)
        
        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404
        
        data = request.get_json()
        
        # Update name
        if 'name' in data and data['name'].strip():
            user.name = data['name'].strip()
            session['user_name'] = user.name  # Update session
        
        # Update email (check for uniqueness)
        if 'email' in data and data['email'].strip():
            new_email = data['email'].strip().lower()
            if new_email != user.email:
                existing = User.get_by_email(new_email)
                if existing:
                    return jsonify({'success': False, 'message': 'Email already exists'}), 400
                user.email = new_email
                session['user_email'] = new_email  # Update session
        
        # Update phone
        if 'phone' in data:
            phone = data['phone'].strip() if data['phone'] else ''
            # Validate phone if provided
            if phone:
                phone = re.sub(r'\D', '', phone)  # Remove non-digits
                if len(phone) != 10:
                    return jsonify({'success': False, 'message': 'Phone must be 10 digits'}), 400
            user.phone = phone if phone else None
        
        # Update password (if provided)
        if 'password' in data and data['password']:
            if len(data['password']) < 8:
                return jsonify({'success': False, 'message': 'Password must be at least 8 characters'}), 400
            user.set_password(data['password'])
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Profile updated successfully',
            'user': {
                'name': user.name,
                'email': user.email,
                'phone': user.phone or ''
            }
        })
    
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error updating profile: {str(e)}')
        return jsonify({'success': False, 'message': str(e)}), 500


# ==================== SUBMIT FOR APPROVAL (USER) ====================
@app.route('/api/drafts/<draft_id>/submit-approval', methods=['POST'])
@login_required
def api_submit_for_approval(draft_id):
    """User submits draft for admin approval"""
    try:
        user_id = session.get('user_id')
        draft = Draft.query.filter_by(id=draft_id, user_id=user_id).first()
        
        if not draft:
            return jsonify({'success': False, 'message': 'Draft not found'}), 404
        
        if draft.status not in ['draft', 'pending']:
            return jsonify({'success': False, 'message': 'Cannot submit this document'}), 400
        
        # Change status to pending
        draft.status = 'pending'
        draft.modified_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Document submitted for approval successfully!'
        })
    
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error submitting for approval: {str(e)}')
        return jsonify({'success': False, 'message': str(e)}), 500


# ==================== CD DOCUMENT PREVIEW (ADMIN ONLY) ====================
@app.route('/api/admin/documents/<doc_id>/cd-preview', methods=['GET'])
@admin_required
def api_get_cd_preview(doc_id):
    """Get CD document preview with filled values for admin"""
    try:
        draft = Draft.query.get(doc_id)
        
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'}), 404
        
        preview_data = draft.preview_data or {}
        template_folder = preview_data.get('template_folder')
        
        if not template_folder:
            # Determine folder from template type
            template_config = TEMPLATE_CONFIG.get(draft.template_type, {})
            folder_type = preview_data.get('folder_type', 'main')
            
            if folder_type == 'unmarried' and 'unmarried_subfolder' in template_config:
                template_folder = template_config['unmarried_subfolder']
            else:
                template_folder = template_config.get('folder', '')
        
        if not template_folder:
            return jsonify({
                'success': False, 
                'message': 'Template folder not configured'
            }), 400
        
        # Get CD content with replacements
        cd_content = get_cd_document_content(template_folder, draft.replacements or {})
        
        return jsonify({
            'success': True,
            'cd_content': cd_content,
            'document_name': draft.old_name or 'Unnamed'
        })
    
    except Exception as e:
        app.logger.error(f'Error getting CD preview: {str(e)}')
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

def get_cd_document_content(template_folder, replacements):
    """
    Get the content of 'CD.docx' document with updated values
    Returns the document content as HTML
    """
    try:
        template_folder_path = Path(template_folder)
        
        # Find CD.docx file specifically (NOT CD_CERTIFICATE.docx or other CD* files)
        cd_file = None
        if template_folder_path.exists():
            for file in template_folder_path.iterdir():
                if file.is_file() and file.suffix.lower() == '.docx':
                    # Check if filename is EXACTLY 'CD.docx' (case-insensitive)
                    if file.stem.upper() == 'CD' and not file.name.startswith('~$'):
                        cd_file = file
                        break
        
        if not cd_file:
            return f"""
                <div class="text-center py-4">
                    <i class="bi bi-file-earmark-x text-muted" style="font-size: 3rem;"></i>
                    <p class="text-muted mt-2">CD.docx document not found in template folder.</p>
                    <small class="text-muted">Looking in: {template_folder}</small>
                </div>
            """
        
        # Open and read the document
        doc = Document(str(cd_file))
        
        # Ensure all replacements are strings (including empty strings for removal)
        clean_replacements = {}
        for key, value in replacements.items():
            if value is not None:
                clean_replacements[key] = str(value).strip()
            else:
                clean_replacements[key] = ''  # Convert None to empty string
        
        # Build HTML content
        html_parts = ['<div class="cd-content-wrapper">']
        
        for para in doc.paragraphs:
            text = para.text
            
            # Skip completely empty paragraphs
            if not text.strip():
                html_parts.append('<p class="cd-paragraph">&nbsp;</p>')
                continue
            
            # Replace all placeholders with actual values (including empty strings)
            for key, value in clean_replacements.items():
                if key in text:
                    if value:  # Non-empty value - highlight it
                        highlighted_value = f'<span class="cd-replaced-value">{value}</span>'
                        text = text.replace(key, highlighted_value)
                    else:  # Empty value - just remove the placeholder
                        text = text.replace(key, '')
            
            # CRITICAL: Clean up extra spaces that result from empty replacements
            # Replace multiple spaces with single space
            text = re.sub(r'\s+', ' ', text)
            
            # Clean up spaces before punctuation
            text = re.sub(r'\s+,', ',', text)
            text = re.sub(r'\s+\.', '.', text)
            text = re.sub(r'\s+;', ';', text)
            
            # Clean up spaces after opening and before closing brackets
            text = re.sub(r'\(\s+', '(', text)
            text = re.sub(r'\s+\)', ')', text)
            
            # Strip the text
            text = text.strip()
            
            # Skip paragraphs that become empty after replacements
            if not text:
                continue
            
            # Determine paragraph style
            style_class = 'cd-paragraph'
            if para.style and para.style.name:
                if 'Heading' in para.style.name:
                    style_class = 'cd-heading'
                elif 'Title' in para.style.name:
                    style_class = 'cd-title'
            
            # Check alignment
            alignment_style = ''
            if para.alignment:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment_style = ' style="text-align: center;"'
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment_style = ' style="text-align: right;"'
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    alignment_style = ' style="text-align: justify;"'
            
            html_parts.append(f'<p class="{style_class}"{alignment_style}>{text}</p>')
        
        # Process tables if any
        for table in doc.tables:
            html_parts.append('<table class="table table-bordered cd-table">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    cell_text = cell.text
                    # Replace placeholders in table cells (including empty strings)
                    for key, value in clean_replacements.items():
                        if key in cell_text:
                            if value:
                                highlighted_value = f'<span class="cd-replaced-value">{value}</span>'
                                cell_text = cell_text.replace(key, highlighted_value)
                            else:
                                cell_text = cell_text.replace(key, '')
                    
                    # Clean up spaces in table cells too
                    cell_text = re.sub(r'\s+', ' ', cell_text).strip()
                    cell_text = re.sub(r'\s+,', ',', cell_text)
                    cell_text = re.sub(r'\s+\.', '.', cell_text)
                    
                    html_parts.append(f'<td>{cell_text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        html_parts.append('</div>')
        
        # Add CSS for highlighting
        css = '''
        <style>
            .cd-content-wrapper {
                font-family: 'Times New Roman', serif;
                font-size: 14px;
                line-height: 1.8;
                color: #333;
                background: white;
                padding: 1.5rem;
                border-radius: 8px;
            }
            .cd-paragraph {
                margin-bottom: 0.75rem;
                text-align: justify;
                line-height: 1.8;
            }
            .cd-heading {
                font-weight: bold;
                font-size: 16px;
                margin-bottom: 1rem;
                text-align: center;
                color: #1a202c;
            }
            .cd-title {
                font-weight: bold;
                font-size: 18px;
                margin-bottom: 1rem;
                text-align: center;
                text-transform: uppercase;
                color: #1a202c;
            }
            .cd-replaced-value {
                background-color: #d4edda;
                color: #155724;
                padding: 2px 6px;
                border-radius: 4px;
                font-weight: 600;
                border: 1px solid #c3e6cb;
            }
            .cd-missing-value {
                background-color: #f8d7da;
                color: #721c24;
                padding: 2px 6px;
                border-radius: 4px;
                font-weight: 600;
                border: 1px solid #f5c6cb;
            }
            .cd-table {
                font-size: 13px;
                margin: 1rem 0;
                width: 100%;
            }
            .cd-table td {
                padding: 0.75rem;
                vertical-align: top;
                border: 1px solid #dee2e6;
            }
            .cd-table tr:hover {
                background-color: #f8f9fa;
            }
        </style>
        '''
        
        return css + ''.join(html_parts)
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"""
            <div class="text-center py-4">
                <i class="bi bi-exclamation-triangle text-danger" style="font-size: 3rem;"></i>
                <p class="text-danger mt-2">Error loading CD document</p>
                <small class="text-muted">{str(e)}</small>
            </div>
        """

# ==================== USER DASHBOARD ROUTES ====================

@app.route('/dashboard')
@login_required
def dashboard():
    """User dashboard with tabs - redirects based on role"""
    user_role = session.get('user_role')
    
    # Redirect super admin to their dashboard
    if user_role == 'super_admin':
        return redirect(url_for('super_admin_dashboard'))
    
    # Redirect admin to their dashboard
    if user_role == 'admin':
        return redirect(url_for('admin_dashboard'))
    
    # Regular user dashboard
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    # Get user stats
    stats = {
        'drafts': Draft.query.filter_by(user_id=user_id, status='draft').count(),
        'pending': Draft.query.filter_by(user_id=user_id, status='pending').count(),
        'approved': Draft.query.filter_by(user_id=user_id, status='approved').count(),
        'generated': Draft.query.filter_by(user_id=user_id, status='generated').count()
    }
    
    return render_template('dashboard.html', 
                          user=user,
                          user_name=user.name,
                          stats=stats)

# ==================== ADMIN DASHBOARD ROUTES ====================

@app.route('/admin/dashboard')
@admin_required
def admin_dashboard():
    """Admin dashboard with tabs"""
    user_id = session.get('user_id')
    admin = User.query.get(user_id)
    
    return render_template('admin_dashboard.html', 
                          admin=admin,
                          admin_name=admin.name)

# ==================== ADMIN API ROUTES ====================

@app.route('/api/admin/stats')
@admin_required
def api_admin_stats():
    """Get admin dashboard statistics"""
    try:
        # User stats
        total_users = User.query.filter_by(role='user').count()
        active_users = User.query.filter_by(role='user', is_active=True).count()
        
        # Draft stats
        total_drafts = Draft.query.count()
        draft_count = Draft.query.filter_by(status='draft').count()
        pending_count = Draft.query.filter_by(status='pending').count()
        approved_count = Draft.query.filter_by(status='approved').count()
        generated_count = Draft.query.filter_by(status='generated').count()
        
        # Per user stats
        users = User.query.filter_by(role='user').all()
        user_stats = []
        for user in users:
            user_stats.append({
                'id': user.id,
                'name': user.name,
                'email': user.email,
                'phone': user.phone,
                'is_active': user.is_active,
                'last_login': user.last_login.isoformat() if user.last_login else None,
                'created_at': user.created_at.isoformat() if user.created_at else None,
                'stats': {
                    'drafts': Draft.query.filter_by(user_id=user.id, status='draft').count(),
                    'pending': Draft.query.filter_by(user_id=user.id, status='pending').count(),
                    'approved': Draft.query.filter_by(user_id=user.id, status='approved').count(),
                    'generated': Draft.query.filter_by(user_id=user.id, status='generated').count(),
                    'total': Draft.query.filter_by(user_id=user.id).count()
                }
            })
        
        return jsonify({
            'success': True,
            'overall': {
                'total_users': total_users,
                'active_users': active_users,
                'total_documents': total_drafts,
                'drafts': draft_count,
                'pending': pending_count,
                'approved': approved_count,
                'generated': generated_count
            },
            'users': user_stats
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/users', methods=['GET'])
@admin_required
def api_admin_get_users():
    """Get all users"""
    try:
        users = User.query.filter_by(role='user').order_by(User.created_at.desc()).all()
        return jsonify({
            'success': True,
            'users': [u.to_dict() for u in users]
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

# ==================== REGISTRATION ROUTE ====================

@app.route('/api/auth/register', methods=['POST'])
def api_register():
    """Register a new user with optional phone verification"""
    try:
        data = request.get_json()
        
        name = data.get('name', '').strip()
        email = data.get('email', '').strip().lower()
        phone = data.get('phone', '').strip()
        password = data.get('password', '').strip()
        phone_verified = data.get('phoneVerified', False)
        
        # Validation
        if not name:
            return jsonify({
                'success': False,
                'message': 'Full name is required'
            }), 400
        
        if not email:
            return jsonify({
                'success': False,
                'message': 'Email is required'
            }), 400
        
        # Validate email format
        if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
            return jsonify({
                'success': False,
                'message': 'Please enter a valid email address'
            }), 400
        
        if not password:
            return jsonify({
                'success': False,
                'message': 'Password is required'
            }), 400
        
        if len(password) < 8:
            return jsonify({
                'success': False,
                'message': 'Password must be at least 8 characters long'
            }), 400
        
        # Validate phone if provided
        if phone:
            phone = re.sub(r'\D', '', phone)  # Remove non-digits
            if len(phone) != 10:
                return jsonify({
                    'success': False,
                    'message': 'Please enter a valid 10-digit phone number'
                }), 400
        
        # Check if email already exists
        existing_user = User.get_by_email(email)
        if existing_user:
            return jsonify({
                'success': False,
                'message': 'An account with this email already exists'
            }), 409
        
        # Check if phone already exists (if provided and verified)
        if phone and phone_verified:
            existing_phone = User.query.filter_by(phone=phone).first()
            if existing_phone:
                return jsonify({
                    'success': False,
                    'message': 'An account with this phone number already exists'
                }), 409
        
        # Create new user
        new_user = User(
            name=name,
            email=email,
            phone=phone if phone else None,
            phone_verified=phone_verified,
            role='user',
            is_active=True
        )
        new_user.set_password(password)
        
        db.session.add(new_user)
        db.session.commit()
        
        app.logger.info(f'New user registered: {email} (Phone verified: {phone_verified})')
        
        return jsonify({
            'success': True,
            'message': 'Account created successfully',
            'user': {
                'id': new_user.id,
                'name': new_user.name,
                'email': new_user.email,
                'phone': new_user.phone,
                'phone_verified': new_user.phone_verified
            }
        }), 201
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Registration Error: {str(e)}')
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': 'An error occurred during registration. Please try again.'
        }), 500

# ==================== OTP VERIFICATION ROUTES ====================

@app.route('/api/otp/send', methods=['POST'])
def api_send_otp():
    """Send OTP to mobile number using KSP API"""
    try:
        data = request.get_json()
        mobile_number = data.get('mobileNumber', '').strip()
        
        app.logger.info(f'OTP Send request for: {mobile_number}')
        
        # Use OTP Service
        result = OTPService.send_otp(mobile_number)
        
        app.logger.info(f'OTP Send result: {result}')
        
        if result['success']:
            return jsonify({
                'success': True,
                'message': result['message']
            }), 200
        else:
            return jsonify({
                'success': False,
                'message': result['message']
            }), 400
            
    except Exception as e:
        app.logger.error(f'OTP Send Error: {str(e)}')
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': 'An error occurred. Please try again.'
        }), 500

@app.route('/api/otp/verify', methods=['POST'])
def api_verify_otp():
    """Verify OTP using KSP API"""
    try:
        data = request.get_json()
        mobile_number = data.get('mobileNumber', '').strip()
        otp = data.get('otp', '').strip()
        
        app.logger.info(f'OTP Verify request for: {mobile_number}, OTP: {otp}')
        
        # Use OTP Service
        result = OTPService.verify_otp(mobile_number, otp)
        
        app.logger.info(f'OTP Verify result: {result}')
        
        # If verified successfully
        if result.get('verified'):
            return jsonify({
                'success': True,
                'message': result['message'],
                'verified': True
            }), 200
        
        # If OTP didn't match but request succeeded
        elif result['success'] and not result.get('verified'):
            return jsonify({
                'success': False,
                'message': result['message'],
                'verified': False
            }), 400
        
        # If request failed
        else:
            return jsonify({
                'success': False,
                'message': result['message'],
                'verified': False
            }), 500
            
    except Exception as e:
        app.logger.error(f'OTP Verify Error: {str(e)}')
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': 'An error occurred. Please try again.'
        }), 500


@app.route('/api/otp/status/<phone>', methods=['GET'])
def api_otp_status(phone):
    """Get OTP status for a phone number"""
    try:
        status = OTPService.get_status(phone)
        return jsonify({
            'success': True,
            'status': status
        }), 200
    except Exception as e:
        app.logger.error(f'OTP Status Error: {str(e)}')
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@app.route('/api/admin/users', methods=['POST'])
@admin_required
def api_admin_create_user():
    """Create a new user"""
    try:
        name = request.json.get('name', '').strip()
        email = request.json.get('email', '').strip().lower()
        phone = request.json.get('phone', '').strip()
        password = request.json.get('password', DEFAULT_USER_PASSWORD)
        role = request.json.get('role', 'user')
        
        if not name or not email:
            return jsonify({'success': False, 'message': 'Name and email are required'})
        
        # Validate email format
        if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
            return jsonify({'success': False, 'message': 'Invalid email format'})
        
        # Check if email exists
        if User.get_by_email(email):
            return jsonify({'success': False, 'message': 'Email already exists'})
        
        # Create user
        new_user = User(
            name=name,
            email=email,
            phone=phone,
            role=role,
            is_active=True
        )
        new_user.set_password(password)
        
        db.session.add(new_user)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'User created successfully',
            'user': new_user.to_dict()
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/users/<user_id>', methods=['PUT'])
@admin_required
def api_admin_update_user(user_id):
    """Update a user"""
    try:
        user = User.query.get(user_id)
        if not user:
            return jsonify({'success': False, 'message': 'User not found'})
        
        name = request.json.get('name', '').strip()
        phone = request.json.get('phone', '').strip()
        new_email = request.json.get('email', '').strip().lower()
        new_password = request.json.get('password', '').strip()
        
        if name:
            user.name = name
        if phone is not None:
            user.phone = phone
        
        if new_email and new_email != user.email:
            existing = User.get_by_email(new_email)
            if existing and existing.id != user_id:
                return jsonify({'success': False, 'message': 'Email already exists'})
            user.email = new_email
        
        if new_password:
            user.set_password(new_password)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'User updated successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/users/<user_id>/toggle', methods=['POST'])
@admin_required
def api_admin_toggle_user(user_id):
    """Toggle user active status"""
    try:
        user = User.query.get(user_id)
        if not user:
            return jsonify({'success': False, 'message': 'User not found'})
        
        user.is_active = not user.is_active
        db.session.commit()
        
        status = 'activated' if user.is_active else 'deactivated'
        return jsonify({'success': True, 'message': f'User {status} successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/users/<user_id>', methods=['DELETE'])
@admin_required
def api_admin_delete_user(user_id):
    """Delete a user"""
    try:
        user = User.query.get(user_id)
        if not user:
            return jsonify({'success': False, 'message': 'User not found'})
        
        db.session.delete(user)
        db.session.commit()
        return jsonify({'success': True, 'message': 'User deleted successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/documents')
@admin_required
def api_admin_get_documents():
    """Get all documents for admin"""
    try:
        status = request.args.get('status')
        
        query = Draft.query.order_by(Draft.modified_at.desc())
        if status:
            query = query.filter_by(status=status)
        
        drafts = query.all()
        return jsonify({
            'success': True,
            'documents': [d.to_dict() for d in drafts]
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

# ==================== DELETE DOCUMENT ENDPOINT ====================
# ==================== DELETE DOCUMENT ENDPOINT ====================
@app.route('/api/admin/documents/<doc_id>', methods=['DELETE'])
@login_required
def delete_document(doc_id):
    """Delete a document (admin only)"""
    try:
        # FIX: Use user_role instead of is_admin
        if session.get('user_role') != 'admin':
            return jsonify({
                'success': False, 
                'message': 'Admin access required'
            }), 403
        
        # Find and delete the draft
        draft = Draft.query.get(doc_id)
        
        if not draft:
            return jsonify({
                'success': False, 
                'message': 'Document not found'
            }), 404
        
        doc_name = draft.old_name or 'Untitled'
        
        # Optional: Delete generated files if they exist
        try:
            if draft.status == 'generated' and draft.preview_data:
                preview_data = draft.preview_data
                output_folder = preview_data.get('output_folder')
                
                if output_folder and os.path.exists(output_folder):
                    import shutil
                    shutil.rmtree(output_folder, ignore_errors=True)
                    print(f"Deleted output folder: {output_folder}")
        except Exception as file_error:
            print(f"File cleanup error (non-critical): {file_error}")
        
        # Delete from database
        db.session.delete(draft)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Document "{doc_name}" deleted successfully'
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting document: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Failed to delete document: {str(e)}'
        }), 500


# ==================== ADMIN DOCUMENT MANAGEMENT ====================

@app.route('/api/admin/documents/<doc_id>', methods=['PUT'])
@admin_required
def api_admin_update_document(doc_id):
    """Admin update document"""
    try:
        draft = Draft.query.get(doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})
        
        data = request.json
        
        if 'replacements' in data:
            draft.replacements = data['replacements']
            draft.old_name = data['replacements'].get('OLD_NAME', draft.old_name)
        
        # Update folder type in preview_data
        folder_type = data.get('folder_type', 'main')
        preview_data = draft.preview_data or {}
        
        # Update template folder based on folder type
        template_type = draft.template_type
        if template_type in TEMPLATE_CONFIG:
            config = TEMPLATE_CONFIG[template_type]
            if folder_type == 'unmarried' and 'unmarried_subfolder' in config:
                preview_data['template_folder'] = config['unmarried_subfolder']
                preview_data['folder_type'] = 'unmarried'
            else:
                preview_data['template_folder'] = config['folder']
                preview_data['folder_type'] = 'main'
        
        draft.preview_data = preview_data
        draft.modified_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Document updated successfully'})
    
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/documents/<doc_id>/approve', methods=['POST'])
@admin_required
def api_admin_approve_document(doc_id):
    """Admin approve document"""
    try:
        draft = Draft.query.get(doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})
        
        draft.status = 'approved'
        draft.approved_at = datetime.utcnow()
        draft.modified_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Document approved successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/documents/<doc_id>/generate', methods=['POST'])
@admin_required
def api_admin_generate_document(doc_id):
    """Admin generate and download document"""
    try:
        draft = Draft.query.get(doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})
        
        if draft.status not in ['approved', 'generated']:
            return jsonify({'success': False, 'message': 'Document must be approved first'})
        
        # Generate documents
        result = generate_documents_to_memory(draft)
        
        if not result['success']:
            return jsonify({'success': False, 'message': result.get('message', 'Generation failed')})
        
        # Update draft status
        draft.status = 'generated'
        draft.generated_at = datetime.utcnow()
        draft.modified_at = datetime.utcnow()
        draft.generated_files = [f"Generated {result['file_count']} files"]
        db.session.commit()
        
        # Return ZIP file
        zip_buffer = result['zip_buffer']
        zip_buffer.seek(0)
        
        filename = f"{draft.old_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/documents/<doc_id>/download', methods=['GET'])
@admin_required
def api_admin_download_document(doc_id):
    """Admin download already generated document"""
    try:
        draft = Draft.query.get(doc_id)
        if not draft:
            return jsonify({'success': False, 'message': 'Document not found'})
        
        # Re-generate from stored data
        result = generate_documents_to_memory(draft)
        if not result['success']:
            return jsonify({'success': False, 'message': result.get('message', 'Generation failed')})
        
        zip_buffer = result['zip_buffer']
        zip_buffer.seek(0)
        filename = f"{draft.old_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/admin/documents/download-bulk', methods=['POST'])
@admin_required
def api_admin_download_bulk():
    """Admin download multiple documents as ZIP"""
    try:
        data = request.json
        doc_ids = data.get('doc_ids', [])
        
        if not doc_ids:
            return jsonify({'success': False, 'message': 'No documents selected'})
        
        # Create a ZIP file in memory
        bulk_zip_buffer = BytesIO()
        
        with zipfile.ZipFile(bulk_zip_buffer, 'w', zipfile.ZIP_DEFLATED) as bulk_zip:
            for doc_id in doc_ids:
                draft = Draft.query.get(doc_id)
                if not draft:
                    continue
                
                # Generate documents for this draft
                result = generate_documents_to_memory(draft)
                if not result['success']:
                    continue
                
                # Add this draft's ZIP to the bulk ZIP
                zip_buffer = result['zip_buffer']
                zip_buffer.seek(0)
                
                # Create a subfolder for each document
                folder_name = draft.old_name.replace(' ', '_') if draft.old_name else f'doc_{doc_id}'
                
                # Extract files from individual ZIP and add to bulk ZIP
                with zipfile.ZipFile(zip_buffer, 'r') as individual_zip:
                    for file_info in individual_zip.filelist:
                        file_data = individual_zip.read(file_info.filename)
                        bulk_zip.writestr(f"{folder_name}/{file_info.filename}", file_data)
        
        bulk_zip_buffer.seek(0)
        
        return send_file(
            bulk_zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f"bulk_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        )
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

def replace_text_in_paragraph(paragraph, replacements):
    """Replace text in paragraph preserving formatting, including empty string replacements"""
    full_text = paragraph.text
    
    # Check if any replacement is needed (including empty string replacements)
    needs_replacement = any(old_text in full_text for old_text in replacements.keys())
    if not needs_replacement or not paragraph.runs:
        return
    
    char_formats = []
    for run in paragraph.runs:
        run_format = {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
            'superscript': run.font.superscript
        }
        for char in run.text:
            char_formats.append(run_format.copy())
    
    combined_text = ''.join(run.text for run in paragraph.runs)
    new_text = combined_text
    new_char_formats = char_formats.copy()
    
    # Sort replacements by length (longer first) to avoid partial replacements
    sorted_replacements = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)
    
    for old_text, new_text_value in sorted_replacements:
        # Handle None values - convert to empty string
        if new_text_value is None:
            new_text_value = ''
        
        # Convert to string
        new_text_value = str(new_text_value)
        
        while old_text in new_text:
            pos = new_text.find(old_text)
            if pos == -1:
                break
            
            if pos < len(new_char_formats):
                placeholder_format = new_char_formats[pos].copy()
            else:
                placeholder_format = char_formats[0].copy() if char_formats else {}
            
            # Replace old text with new text (can be empty string)
            new_text = new_text[:pos] + new_text_value + new_text[pos + len(old_text):]
            
            # Remove format entries for old text
            del new_char_formats[pos:pos + len(old_text)]
            
            # Add format entries for new text (if not empty)
            for i in range(len(new_text_value)):
                format_copy = placeholder_format.copy()
                
                if old_text == "ALPHA_DATE":
                    match = re.match(r'(\d{1,2})(ST|ND|RD|TH)', new_text_value, re.IGNORECASE)
                    if match:
                        day_len = len(match.group(1))
                        suffix_len = len(match.group(2))
                        if i >= day_len and i < day_len + suffix_len:
                            format_copy['superscript'] = True
                
                new_char_formats.insert(pos + i, format_copy)
    
    # CRITICAL: Clean up extra spaces that result from empty replacements
    # Replace multiple spaces with single space
    original_len = len(new_text)
    while '  ' in new_text:
        double_space_pos = new_text.find('  ')
        new_text = new_text[:double_space_pos] + ' ' + new_text[double_space_pos + 2:]
        # Also remove corresponding format entry
        if double_space_pos < len(new_char_formats):
            del new_char_formats[double_space_pos]
    
    # Clean up spaces before commas (e.g., " ," becomes ",")
    while ' ,' in new_text:
        pos = new_text.find(' ,')
        new_text = new_text[:pos] + ',' + new_text[pos + 2:]
        if pos < len(new_char_formats):
            del new_char_formats[pos]
    
    # Clean up spaces before periods (e.g., " ." becomes ".")
    while ' .' in new_text:
        pos = new_text.find(' .')
        new_text = new_text[:pos] + '.' + new_text[pos + 2:]
        if pos < len(new_char_formats):
            del new_char_formats[pos]
    
    # Strip leading/trailing spaces from the final text
    leading_spaces = len(new_text) - len(new_text.lstrip())
    trailing_spaces = len(new_text) - len(new_text.rstrip())
    
    if leading_spaces > 0:
        new_text = new_text[leading_spaces:]
        new_char_formats = new_char_formats[leading_spaces:]
    
    if trailing_spaces > 0 and len(new_text) > 0:
        new_text = new_text[:-trailing_spaces] if trailing_spaces > 0 else new_text
        new_char_formats = new_char_formats[:-trailing_spaces] if trailing_spaces > 0 else new_char_formats
    
    # Clear all runs
    for run in paragraph.runs:
        run.text = ''
    
    if not new_char_formats or not new_text:
        if new_text:
            paragraph.runs[0].text = new_text
        return
    
    def format_key(fmt):
        if fmt is None:
            return None
        return (fmt.get('bold'), fmt.get('italic'), fmt.get('underline'),
                fmt.get('font_name'), fmt.get('font_size'), str(fmt.get('font_color')), fmt.get('superscript'))
    
    groups = []
    current_group = {'text': '', 'format': new_char_formats[0] if new_char_formats else None}
    
    for i, char in enumerate(new_text):
        char_format = new_char_formats[i] if i < len(new_char_formats) else (new_char_formats[-1] if new_char_formats else None)
        
        if format_key(char_format) == format_key(current_group['format']):
            current_group['text'] += char
        else:
            if current_group['text']:
                groups.append(current_group)
            current_group = {'text': char, 'format': char_format}
    
    if current_group['text']:
        groups.append(current_group)
    
    if groups:
        paragraph.runs[0].text = groups[0]['text']
        apply_format(paragraph.runs[0], groups[0]['format'])
        
        for group in groups[1:]:
            new_run = paragraph.add_run(group['text'])
            apply_format(new_run, group['format'])


# ==================== HELPER: GENERATE TO MEMORY ====================
def generate_documents_to_memory(draft):
    """Generate documents in memory and return as ZIP buffer"""
    try:
        template_type = draft.template_type
        replacements = draft.replacements or {}
        preview_data = draft.preview_data or {}
        
        if template_type not in TEMPLATE_CONFIG:
            return {'success': False, 'message': 'Invalid template type'}
        
        template_config = TEMPLATE_CONFIG[template_type]
        
        # Get template folder from preview_data or determine from relation
        template_folder_str = preview_data.get('template_folder')
        
        if not template_folder_str:
            # Determine folder based on relation
            relation = replacements.get('UPDATE_RELATION', '')
            folder_type = preview_data.get('folder_type', 'main')
            
            if folder_type == 'unmarried' or (relation == 'D/o' and not replacements.get('SPOUSE_NAME1')):
                if 'unmarried_subfolder' in template_config:
                    template_folder_str = template_config['unmarried_subfolder']
                else:
                    template_folder_str = template_config['folder']
            else:
                template_folder_str = template_config['folder']
        
        template_folder = Path(template_folder_str)
        
        templates = get_templates(template_folder)
        if not templates:
            return {'success': False, 'message': f'No templates found in {template_folder}'}
        
        # Determine the folder name source based on template type
        # CRITICAL FIX: For minor template, use FATHER-MOTHER_NAME instead of OLD_NAME
        if template_type == 'minor_template':
            folder_name_source = replacements.get('FATHER-MOTHER_NAME', '')
            if not folder_name_source:
                folder_name_source = replacements.get('OLD_NAME', 'unnamed')
        else:
            folder_name_source = replacements.get('OLD_NAME', 'unnamed')
        
        # Create ZIP file in memory
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for template_name in templates:
                input_path = template_folder / template_name
                
                # Generate document in memory
                doc = Document(str(input_path))
                
                # Replace text (including empty strings for WIFE_OF, SPOUSE_NAME1, etc.)
                for paragraph in doc.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
                
                if doc.tables:
                    replace_text_in_tables(doc.tables, replacements)
                
                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    if header.tables:
                        replace_text_in_tables(header.tables, replacements)
                
                for section in doc.sections:
                    footer = section.footer
                    for paragraph in footer.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    if footer.tables:
                        replace_text_in_tables(footer.tables, replacements)
                
                # Save to memory buffer
                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                # Add to ZIP - use the correct folder name source
                original_name = Path(template_name).stem
                safe_folder = create_safe_folder_name(folder_name_source)
                output_filename = f"{original_name} {safe_folder}.docx"
                
                zipf.writestr(output_filename, doc_buffer.read())
        
        return {
            'success': True,
            'zip_buffer': zip_buffer,
            'file_count': len(templates)
        }
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {'success': False, 'message': str(e)}
    
# ==================== DRAFT API ROUTES ====================

@app.route('/api/drafts', methods=['GET'])
@login_required
def api_get_drafts():
    """Get user's drafts"""
    try:
        user_id = session.get('user_id')
        status = request.args.get('status')
        
        query = Draft.query.filter_by(user_id=user_id).order_by(Draft.modified_at.desc())
        if status:
            query = query.filter_by(status=status)
        
        drafts = query.all()
        return jsonify({
            'success': True,
            'drafts': [d.to_dict() for d in drafts]
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/drafts', methods=['POST'])
@login_required
def api_create_draft():
    """Create a new draft"""
    try:
        user_id = session.get('user_id')
        data = request.json
        
        template_type = data.get('template_type')
        template_config = TEMPLATE_CONFIG.get(template_type, {})
        
        new_draft = Draft(
            user_id=user_id,
            template_type=template_type,
            template_name=template_config.get('name', template_type),
            old_name=data.get('replacements', {}).get('OLD_NAME', ''),
            replacements=data.get('replacements', {}),
            preview_data=data.get('preview_data', {}),
            status='draft'
        )
        
        db.session.add(new_draft)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Draft saved successfully',
            'draft': new_draft.to_dict()
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/drafts/<draft_id>', methods=['GET'])
@login_required
def api_get_draft(draft_id):
    """Get a specific draft"""
    try:
        user_id = session.get('user_id')
        draft = Draft.query.filter_by(id=draft_id, user_id=user_id).first()
        
        if not draft:
            return jsonify({'success': False, 'message': 'Draft not found'})
        
        return jsonify({'success': True, 'draft': draft.to_dict()})
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/drafts/<draft_id>', methods=['PUT'])
@login_required
def api_update_draft(draft_id):
    """Update a draft"""
    try:
        user_id = session.get('user_id')
        draft = Draft.query.filter_by(id=draft_id, user_id=user_id).first()
        
        if not draft:
            return jsonify({'success': False, 'message': 'Draft not found'})
        
        data = request.json
        
        if 'replacements' in data:
            draft.replacements = data['replacements']
            draft.old_name = data['replacements'].get('OLD_NAME', draft.old_name)
        
        if 'preview_data' in data:
            draft.preview_data = data['preview_data']
        
        # Handle folder_type update
        if 'folder_type' in data:
            preview_data = draft.preview_data or {}
            template_type = draft.template_type
            if template_type in TEMPLATE_CONFIG:
                config = TEMPLATE_CONFIG[template_type]
                if data['folder_type'] == 'unmarried' and 'unmarried_subfolder' in config:
                    preview_data['template_folder'] = config['unmarried_subfolder']
                    preview_data['folder_type'] = 'unmarried'
                else:
                    preview_data['template_folder'] = config['folder']
                    preview_data['folder_type'] = 'main'
            draft.preview_data = preview_data
        
        draft.modified_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Draft updated successfully', 'draft': draft.to_dict()})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/drafts/<draft_id>/approve', methods=['POST'])
@login_required
def api_approve_draft(draft_id):
    """Approve a draft"""
    try:
        user_id = session.get('user_id')
        draft = Draft.query.filter_by(id=draft_id, user_id=user_id).first()
        
        if not draft:
            return jsonify({'success': False, 'message': 'Draft not found'})
        
        draft.status = 'approved'
        draft.approved_at = datetime.utcnow()
        draft.modified_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Draft approved successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/drafts/<draft_id>', methods=['DELETE'])
@login_required
def api_delete_draft(draft_id):
    """Delete a draft"""
    try:
        user_id = session.get('user_id')
        draft = Draft.query.filter_by(id=draft_id, user_id=user_id).first()
        
        if not draft:
            return jsonify({'success': False, 'message': 'Draft not found'})
        
        db.session.delete(draft)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Draft deleted successfully'})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/drafts/stats')
@login_required
def api_draft_stats():
    """Get draft statistics for current user"""
    try:
        user_id = session.get('user_id')
        
        stats = {
            'drafts': Draft.query.filter_by(user_id=user_id, status='draft').count(),
            'pending': Draft.query.filter_by(user_id=user_id, status='pending').count(),
            'approved': Draft.query.filter_by(user_id=user_id, status='approved').count(),
            'generated': Draft.query.filter_by(user_id=user_id, status='generated').count()
        }
        stats['total'] = sum(stats.values())
        
        return jsonify({'success': True, 'stats': stats})
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

# ==================== GENERATE FROM APPROVED DRAFTS ====================

@app.route('/api/generate/approved', methods=['GET'])
@login_required
def api_get_approved():
    """Get approved drafts ready for generation"""
    try:
        user_id = session.get('user_id')
        
        approved = Draft.query.filter_by(user_id=user_id, status='approved').order_by(Draft.approved_at.desc()).all()
        generated = Draft.query.filter_by(user_id=user_id, status='generated').order_by(Draft.generated_at.desc()).all()
        
        return jsonify({
            'success': True,
            'approved': [d.to_dict() for d in approved],
            'generated': [d.to_dict() for d in generated]
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/generate/batch', methods=['POST'])
@login_required
def api_generate_batch():
    """Generate documents for multiple approved drafts and return as ZIP"""
    try:
        user_id = session.get('user_id')
        draft_ids = request.json.get('draft_ids', [])
        
        if not draft_ids:
            return jsonify({'success': False, 'message': 'No drafts selected'})
        
        generated_files = []
        errors = []
        
        # Create bulk ZIP
        bulk_zip_buffer = BytesIO()
        
        with zipfile.ZipFile(bulk_zip_buffer, 'w', zipfile.ZIP_DEFLATED) as bulk_zip:
            for draft_id in draft_ids:
                draft = Draft.query.filter_by(id=draft_id, user_id=user_id, status='approved').first()
                
                if not draft:
                    errors.append({'draft_id': draft_id, 'error': 'Draft not found or not approved'})
                    continue
                
                try:
                    result = generate_documents_to_memory(draft)
                    
                    if result['success']:
                        # Update draft status
                        draft.status = 'generated'
                        draft.generated_at = datetime.utcnow()
                        draft.generated_files = [f"Generated {result['file_count']} files"]
                        
                        # Add to bulk ZIP
                        zip_buffer = result['zip_buffer']
                        zip_buffer.seek(0)
                        
                        folder_name = draft.old_name.replace(' ', '_') if draft.old_name else f'doc_{draft_id}'
                        
                        with zipfile.ZipFile(zip_buffer, 'r') as individual_zip:
                            for file_info in individual_zip.filelist:
                                file_data = individual_zip.read(file_info.filename)
                                bulk_zip.writestr(f"{folder_name}/{file_info.filename}", file_data)
                        
                        generated_files.append({
                            'draft_id': draft_id,
                            'old_name': draft.old_name,
                            'file_count': result['file_count']
                        })
                    else:
                        errors.append({'draft_id': draft_id, 'error': result.get('message')})
                except Exception as e:
                    errors.append({'draft_id': draft_id, 'error': str(e)})
        
        db.session.commit()
        
        if generated_files:
            bulk_zip_buffer.seek(0)
            return send_file(
                bulk_zip_buffer,
                mimetype='application/zip',
                as_attachment=True,
                download_name=f"documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
            )
        else:
            return jsonify({
                'success': False,
                'message': 'No documents were generated',
                'errors': errors
            })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/generate/single/<draft_id>', methods=['POST'])
@login_required
def api_generate_single(draft_id):
    """Generate a single draft and download"""
    try:
        user_id = session.get('user_id')
        draft = Draft.query.filter_by(id=draft_id, user_id=user_id, status='approved').first()
        
        if not draft:
            return jsonify({'success': False, 'message': 'Draft not found or not approved'})
        
        result = generate_documents_to_memory(draft)
        
        if not result['success']:
            return jsonify({'success': False, 'message': result.get('message', 'Generation failed')})
        
        # Update draft status
        draft.status = 'generated'
        draft.generated_at = datetime.utcnow()
        draft.generated_files = [f"Generated {result['file_count']} files"]
        db.session.commit()
        
        # Return ZIP file
        zip_buffer = result['zip_buffer']
        zip_buffer.seek(0)
        
        filename = f"{draft.old_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)})

# ==================== PHONE NUMBER MANAGEMENT ====================

def load_phone_numbers_from_csv():
    """Load phone numbers from CSV"""
    phone_numbers = []
    if os.path.exists(PHONE_CSV_FILE):
        try:
            with open(PHONE_CSV_FILE, 'r', newline='', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    if row and row[0].strip():
                        phone = re.sub(r'\D', '', row[0].strip())
                        if len(phone) == 10:
                            phone_numbers.append(phone)
        except Exception as e:
            print(f"Error loading phones: {e}")
    return phone_numbers

def get_session_reserved_phones():
    if 'reserved_phones' not in session:
        session['reserved_phones'] = []
    return session['reserved_phones']

def add_session_reserved_phone(phone):
    reserved = get_session_reserved_phones()
    if phone not in reserved:
        reserved.append(phone)
        session['reserved_phones'] = reserved
        session.modified = True

def clear_session_reserved_phones():
    session['reserved_phones'] = []
    session.modified = True

def get_next_available_phone(exclude_phones=None):
    """Get next available phone number from CSV, ensuring no repetition until all are used"""
    all_phones = load_phone_numbers_from_csv()
    
    if not all_phones:
        return None, "No phone numbers found in CSV"
    
    # Get used phones from database
    used_phones = {pt.phone for pt in PhoneTracking.query.filter_by(is_used=True).all()}
    
    # Get session reserved phones (phones assigned in current form session)
    session_reserved = set(get_session_reserved_phones())
    
    # Get phones to exclude (already selected in the current form)
    exclude_set = set(exclude_phones) if exclude_phones else set()
    
    # Combine all excluded phones
    all_excluded = used_phones | session_reserved | exclude_set
    
    # Find available phones
    available_phones = [p for p in all_phones if p not in all_excluded]
    
    # If no phones available, reset the database tracking and try again
    if not available_phones:
        # Check if there are phones available after excluding only session and form phones
        potential_phones = [p for p in all_phones if p not in session_reserved and p not in exclude_set]
        
        if potential_phones:
            # Reset database tracking - all phones can be used again
            PhoneTracking.query.update({PhoneTracking.is_used: False})
            db.session.commit()
            available_phones = potential_phones
            print(f"Phone tracking reset. {len(available_phones)} phones now available.")
        else:
            # All phones are either in session or excluded - cannot proceed
            return None, "All phone numbers are in use. Please complete current form or start fresh."
    
    if not available_phones:
        return None, "No phone numbers available"
    
    # Get the next phone (first available)
    next_phone = available_phones[0]
    
    # Add to session reserved (so it won't be assigned again in this session)
    add_session_reserved_phone(next_phone)
    
    return next_phone, None

def mark_phones_as_used(phone_numbers):
    if not phone_numbers:
        return
    
    for phone in phone_numbers:
        if phone and isinstance(phone, str):
            phone = re.sub(r'\D', '', phone.strip())
            if len(phone) == 10:
                existing = PhoneTracking.query.filter_by(phone=phone).first()
                if existing:
                    existing.is_used = True
                    existing.used_at = datetime.utcnow()
                else:
                    pt = PhoneTracking(phone=phone, is_used=True, used_at=datetime.utcnow())
                    db.session.add(pt)
    
    db.session.commit()

def release_session_phone(phone):
    reserved = get_session_reserved_phones()
    if phone in reserved:
        reserved.remove(phone)
        session['reserved_phones'] = reserved
        session.modified = True

def get_phone_stats():
    all_phones = load_phone_numbers_from_csv()
    used_count = PhoneTracking.query.filter_by(is_used=True).count()
    session_reserved = len(get_session_reserved_phones())
    available = len(all_phones) - used_count - session_reserved
    
    return {
        'total': len(all_phones),
        'used': used_count,
        'reserved': session_reserved,
        'available': max(0, available)
    }

@app.route('/api/phone/next', methods=['POST'])
def api_get_next_phone():
    data = request.get_json() or {}
    exclude_phones = data.get('exclude', [])
    
    phone, error = get_next_available_phone(exclude_phones)
    
    if error:
        return jsonify({'success': False, 'message': error, 'stats': get_phone_stats()})
    
    return jsonify({'success': True, 'phone': phone, 'stats': get_phone_stats()})

@app.route('/api/phone/release', methods=['POST'])
def api_release_phone():
    phone = request.json.get('phone', '')
    phone = re.sub(r'\D', '', phone.strip())
    
    if len(phone) == 10:
        release_session_phone(phone)
    
    return jsonify({'success': True, 'message': 'Phone released', 'stats': get_phone_stats()})

@app.route('/api/phone/stats')
def api_phone_stats():
    return jsonify({'success': True, 'stats': get_phone_stats()})

@app.route('/api/phone/reset', methods=['POST'])
def api_reset_phone_tracking():
    PhoneTracking.query.update({PhoneTracking.is_used: False})
    db.session.commit()
    clear_session_reserved_phones()
    return jsonify({'success': True, 'message': 'Phone tracking reset', 'stats': get_phone_stats()})

@app.route('/api/phone/clear_session', methods=['POST'])
@login_required
def api_clear_session_phones():
    """Clear session reserved phones"""
    try:
        clear_session_reserved_phones()
        return jsonify({'success': True, 'message': 'Session phones cleared', 'stats': get_phone_stats()})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

# ==================== HELPER FUNCTIONS ====================

def to_uppercase(value):
    if isinstance(value, str):
        return value.upper().strip()
    return value

def to_uppercase_preserve_alias(value):
    if not isinstance(value, str):
        return value
    
    value = value.strip()
    if not value:
        return value
    
    parts = re.split(r'\s+alias\s+', value, flags=re.IGNORECASE)
    uppercased_parts = [part.strip().upper() for part in parts if part.strip()]
    
    return ' alias '.join(uppercased_parts)

def format_date_to_ddmmyyyy(date_string):
    if not date_string:
        return ''
    try:
        date_obj = datetime.strptime(date_string, '%Y-%m-%d')
        return date_obj.strftime('%d/%m/%Y')
    except ValueError:
        return date_string

def generate_email_from_name(name):
    if not name:
        return ''
    
    if ' alias ' in name.lower():
        name = re.split(r'\s+alias\s+', name, flags=re.IGNORECASE)[0].strip()
    
    clean_name = re.sub(r'[^a-zA-Z]', '', name).lower()
    
    if not clean_name:
        return ''
    
    random_digits = str(random.randint(10, 999))
    return f"{clean_name}{random_digits}@gmail.com"

def get_gender_pronouns(son_daughter=None, gender=None):
    pronouns = {}
    
    if son_daughter:
        son_daughter_lower = son_daughter.lower().strip()
        if son_daughter_lower in GENDER_PRONOUNS:
            pronouns = GENDER_PRONOUNS[son_daughter_lower].copy()
    elif gender:
        gender_lower = gender.lower().strip()
        if gender_lower in GENDER_PRONOUNS_BY_GENDER:
            pronouns = GENDER_PRONOUNS_BY_GENDER[gender_lower].copy()
    
    return pronouns

def create_safe_folder_name(name):
    if not name:
        return 'unnamed'
    if ' alias ' in name.lower():
        name = re.split(r'\s+alias\s+', name, flags=re.IGNORECASE)[0].strip()
    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()
    return safe_name if safe_name else 'unnamed'


def apply_format(run, formatting):
    if formatting is None:
        return
    
    if formatting.get('bold') is not None:
        run.bold = formatting['bold']
    if formatting.get('italic') is not None:
        run.italic = formatting['italic']
    if formatting.get('underline') is not None:
        run.underline = formatting['underline']
    if formatting.get('font_name'):
        run.font.name = formatting['font_name']
    if formatting.get('font_size'):
        run.font.size = formatting['font_size']
    if formatting.get('font_color'):
        run.font.color.rgb = formatting['font_color']
    if formatting.get('superscript') is not None:
        run.font.superscript = formatting['superscript']

def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

def get_templates(folder):
    templates = []
    folder_path = Path(folder)
    if folder_path.exists():
        for file in folder_path.iterdir():
            if file.is_file() and file.suffix.lower() == '.docx' and not file.name.startswith('~$'):
                templates.append(file.name)
    return sorted(templates)

def get_all_template_info():
    template_info = {}
    for key, config in TEMPLATE_CONFIG.items():
        folder = config['folder']
        templates = get_templates(folder)
        template_info[key] = {
            'name': config['name'],
            'description': config['description'],
            'folder': folder,
            'icon': config['icon'],
            'color': config['color'],
            'count': len(templates),
            'files': templates
        }
    return template_info

# ==================== MAIN ROUTES ====================

@app.route('/')
@app.route('/home')
def index():
    """Landing page - always accessible, no login required"""
    return render_template('index.html')

@app.route('/namechangeservice')
@login_required
def namechangeservice():
    clear_session_reserved_phones()
    
    template_info = get_all_template_info()
    phone_stats = get_phone_stats()
    
    return render_template('namechangeservice.html',
                         template_info=template_info,
                         template_config=TEMPLATE_CONFIG,
                         cast_options=CAST_OPTIONS,
                         phone_stats=phone_stats,
                         user_name=session.get('user_name'))

@app.route('/get_template_config/<template_type>')
@login_required
def get_template_config_route(template_type):
    if template_type in TEMPLATE_CONFIG:
        config = TEMPLATE_CONFIG[template_type]
        templates = get_templates(config['folder'])
        return jsonify({
            'success': True,
            'config': config,
            'templates': templates,
            'count': len(templates),
            'folder': config['folder']
        })
    return jsonify({'success': False, 'message': 'Template type not found'})

@app.route('/get_templates_by_relation/<template_type>/<relation>')
@login_required
def get_templates_by_relation(template_type, relation):
    if template_type not in TEMPLATE_CONFIG:
        return jsonify({'success': False, 'message': 'Template type not found'})
    
    config = TEMPLATE_CONFIG[template_type]
    
    if template_type in ['major_template', 'religion_template'] and relation.lower() == 'd':
        if 'unmarried_subfolder' in config:
            folder = config['unmarried_subfolder']
            templates = get_templates(folder)
            return jsonify({
                'success': True,
                'templates': templates,
                'count': len(templates),
                'folder': folder,
                'folder_type': 'unmarried'
            })
    
    folder = config['folder']
    templates = get_templates(folder)
    return jsonify({
        'success': True,
        'templates': templates,
        'count': len(templates),
        'folder': folder,
        'folder_type': 'main'
    })

@app.route('/generate_email', methods=['POST'])
@login_required
def generate_email_route():
    name = request.json.get('name', '')
    email = generate_email_from_name(name)
    return jsonify({'success': True, 'email': email})

@app.route('/preview', methods=['POST'])
@login_required
def preview_document():
    try:
        template_type = request.form.get('template_type', '').strip()
        
        if template_type not in TEMPLATE_CONFIG:
            return jsonify({'success': False, 'message': 'Invalid template type!'})
        
        template_config = TEMPLATE_CONFIG[template_type]
        
        relation_input = request.form.get('relation', '').strip().lower()
        update_relation = RELATION_MAPPING.get(relation_input, '')
        
        # Determine template folder based on relation and template type
        if template_type in ['major_template', 'religion_template'] and relation_input == 'd':
            # Check if spouse name is provided for D/o relation
            spouse_name_check = request.form.get('spouse_name', '').strip() if relation_input == 'd/w' else ''
            
            # Use unmarried folder if D/o without spouse
            if not spouse_name_check and 'unmarried_subfolder' in template_config:
                template_folder = template_config['unmarried_subfolder']
                folder_type = 'unmarried'
            else:
                template_folder = template_config['folder']
                folder_type = 'main'
        else:
            template_folder = template_config['folder']
            folder_type = 'main'
        
        templates = get_templates(template_folder)
        
        if not templates:
            return jsonify({'success': False, 'message': f'No template files found!'})
        
        old_name_raw = request.form.get('old_name', '').strip()
        old_name = to_uppercase_preserve_alias(old_name_raw)
        
        new_name = to_uppercase(request.form.get('new_name', ''))
        gender = to_uppercase(request.form.get('gender_update', ''))
        
        if not old_name:
            return jsonify({'success': False, 'message': 'Name field is required!'})
        
        base_old_name = re.split(r'\s+alias\s+', old_name, flags=re.IGNORECASE)[0].strip()
        
        phone_update = request.form.get('phone_update', '').strip()
        witness_phone1 = request.form.get('witness_phone1', '').strip()
        witness_phone2 = request.form.get('witness_phone2', '').strip()
        
        phones_used = [p for p in [phone_update, witness_phone1, witness_phone2] if p]
        
        # Initialize base replacements
        replacements = {
            "OLD_NAME": old_name,
            "NEW_NAME": new_name if new_name else base_old_name,
            "UPDATE_ADDRESS": to_uppercase(request.form.get('update_address', '')),
            "GENDER_UPDATE": gender,
            "CAST_UPDATE": to_uppercase(request.form.get('cast_update', '')),
            "PHONE_UPDATE": phone_update,
            "EMAIL_UPDATE": request.form.get('email_update', '').strip(),
            "NUM_DATE": request.form.get('num_date', '').strip(),
            "ALPHA_DATE": request.form.get('alpha_date', '').strip(),
            "WITNESS_NAME1": to_uppercase(request.form.get('witness_name1', '')),
            "WITNESS_ADDRESS1": to_uppercase(request.form.get('witness_address1', '')),
            "WITNESS_PHONE1": witness_phone1,
            "WITNESS_NAME2": to_uppercase(request.form.get('witness_name2', '')),
            "WITNESS_ADDRESS2": to_uppercase(request.form.get('witness_address2', '')),
            "WITNESS_PHONE2": witness_phone2,
        }
        
        # Handle relation-specific fields - CRITICAL: Set WIFE_OF and SPOUSE_NAME1 properly
        if relation_input == 'd/w':
            # D/o & W/o selected - use both father and spouse names
            father_name = to_uppercase(request.form.get('father_name', ''))
            spouse_name = to_uppercase(request.form.get('spouse_name', ''))

            replacements.update({
                "UPDATE_RELATION": "D/o",
                "FATHER-SPOUSE_NAME": father_name,
                "WIFE_OF": " W/o ",
                "SPOUSE_NAME1": spouse_name,
            })
        else:
            # NOT D/o & W/o - WIFE_OF and SPOUSE_NAME1 MUST be empty strings
            fatherspouse_name = to_uppercase(request.form.get('fatherspouse_name', ''))
            replacements.update({
                "UPDATE_RELATION": update_relation,
                "FATHER-SPOUSE_NAME": fatherspouse_name,
                "WIFE_OF": "",  # Empty string - will be removed from document
                "SPOUSE_NAME1": "",  # Empty string - will be removed from document
            })
        
        son_daughter = None
        folder_name_source = base_old_name
        
        # Handle minor template specific fields
        if template_type == 'minor_template':
            child_dob_raw = request.form.get('child_dob', '').strip()
            child_dob_formatted = format_date_to_ddmmyyyy(child_dob_raw)
            
            son_daughter = request.form.get('son_daughter', '').strip()
            fathermother_name = to_uppercase(request.form.get('fathermother_name', ''))
            
            # Handle guardian relation for minor
            if relation_input == 'd/w':
                guardian_father_name = to_uppercase(request.form.get('guardian_father_name', ''))
                guardian_spouse_name = to_uppercase(request.form.get('guardian_spouse_name', ''))
                
                replacements.update({
                    "UPDATE_RELATION": "D/o",
                    "FATHER-SPOUSE_NAME": guardian_father_name,
                    "WIFE_OF": " W/o ",
                    "SPOUSE_NAME1": guardian_spouse_name,
                })
            
            replacements.update({
                "UPDATE_AGE": request.form.get('update_age', '').strip(),
                "FATHER-MOTHER_NAME": fathermother_name,
                "SON-DAUGHTER": son_daughter,
                "CHILD_DOB": child_dob_formatted,
                "BIRTH_PLACE": to_uppercase(request.form.get('birth_place', '')),
            })
            
            # CRITICAL FIX: For minor template, use Father/Mother Name for folder name
            if fathermother_name:
                folder_name_source = fathermother_name
            else:
                folder_name_source = base_old_name
        
        # Get gender pronouns and add to replacements
        gender_pronouns = get_gender_pronouns(son_daughter=son_daughter, gender=gender)
        replacements.update(gender_pronouns)
        
        # Ensure HE_SHE is set (critical fix)
        if 'HE_SHE' not in replacements or not replacements['HE_SHE']:
            if son_daughter:
                replacements['HE_SHE'] = 'he' if son_daughter.lower() == 'son' else 'she'
            elif gender:
                replacements['HE_SHE'] = 'he' if gender.lower() == 'male' else 'she' if gender.lower() == 'female' else 'he/she'
            else:
                replacements['HE_SHE'] = 'he/she'  # Default fallback
        
        # Store preview data in session
        session['preview_data'] = {
            'template_type': template_type,
            'template_folder': str(template_folder),
            'templates': templates,
            'replacements': replacements,
            'folder_name_source': folder_name_source,
            'relation_input': relation_input,
            'folder_type': folder_type,
            'phones_used': phones_used
        }
        
        # Clean display replacements (remove empty values for display, but keep them for processing)
        display_replacements = {k: v for k, v in replacements.items() 
                               if v and str(v).strip()}
        
        return jsonify({
            'success': True,
            'message': 'Preview generated successfully!',
            'template_type': template_type,
            'template_name': template_config['name'],
            'template_count': len(templates),
            'replacements': display_replacements,
            'folder_named_by': 'Father/Mother Name' if template_type == 'minor_template' else 'Applicant Name',
            'used_unmarried_folder': folder_type == 'unmarried'
        })
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

@app.route('/update_preview', methods=['POST'])
@login_required
def update_preview():
    try:
        if 'preview_data' not in session:
            return jsonify({'success': False, 'message': 'No preview data found.'})
        
        updated_replacements = request.json.get('replacements', {})
        
        for key, value in updated_replacements.items():
            if key in session['preview_data']['replacements']:
                session['preview_data']['replacements'][key] = value
        
        session.modified = True
        
        return jsonify({'success': True, 'message': 'Preview updated!'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

@app.route('/generate', methods=['POST'])
@login_required
def generate_document():
    """Generate and download documents directly"""
    try:
        if 'preview_data' not in session:
            return jsonify({'success': False, 'message': 'No preview data found.'})
        
        preview_data = session['preview_data']
        
        template_type = preview_data['template_type']
        template_folder = Path(preview_data['template_folder'])
        templates = preview_data['templates']
        replacements = preview_data['replacements']
        
        phones_used = preview_data.get('phones_used', [])
        mark_phones_as_used(phones_used)
        clear_session_reserved_phones()
        
        # Determine the folder name source based on template type
        # CRITICAL FIX: For minor template, use Father/Mother Name
        if template_type == 'minor_template':
            folder_name_source = replacements.get('FATHER-MOTHER_NAME', '')
            if not folder_name_source:
                folder_name_source = preview_data.get('folder_name_source', replacements.get('OLD_NAME', 'unnamed'))
        else:
            folder_name_source = preview_data.get('folder_name_source', replacements.get('OLD_NAME', 'unnamed'))
        
        # Generate documents in memory
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for template_name in templates:
                input_path = template_folder / template_name
                
                doc = Document(str(input_path))
                
                for paragraph in doc.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
                
                if doc.tables:
                    replace_text_in_tables(doc.tables, replacements)
                
                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    if header.tables:
                        replace_text_in_tables(header.tables, replacements)
                
                for section in doc.sections:
                    footer = section.footer
                    for paragraph in footer.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    if footer.tables:
                        replace_text_in_tables(footer.tables, replacements)
                
                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                original_name = Path(template_name).stem
                safe_folder = create_safe_folder_name(folder_name_source)
                output_filename = f"{original_name} {safe_folder}.docx"
                
                zipf.writestr(output_filename, doc_buffer.read())
        
        session.pop('preview_data', None)
        
        zip_buffer.seek(0)
        filename = f"{create_safe_folder_name(folder_name_source)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})
    

@app.route('/api/drafts/save', methods=['POST'])
@login_required
def api_save_draft_from_preview():
    """Save a new draft from preview data or direct submission"""
    try:
        user_id = session.get('user_id')
        data = request.json
        
        # Get data from request or session
        replacements = data.get('replacements', {})
        template_type = data.get('template_type', '')
        preview_data = data.get('preview_data', {})
        
        # If we have session preview_data, use that
        if 'preview_data' in session and not replacements:
            session_preview = session['preview_data']
            replacements = session_preview.get('replacements', {})
            template_type = session_preview.get('template_type', template_type)
            preview_data = session_preview
        
        if not template_type:
            return jsonify({'success': False, 'message': 'Template type is required'})
        
        if not replacements:
            return jsonify({'success': False, 'message': 'No data to save'})
        
        template_config = TEMPLATE_CONFIG.get(template_type, {})
        
        # Create new draft
        new_draft = Draft(
            user_id=user_id,
            template_type=template_type,
            template_name=template_config.get('name', template_type),
            old_name=replacements.get('OLD_NAME', 'Unnamed'),
            replacements=replacements,
            preview_data=preview_data,
            status='draft'
        )
        
        db.session.add(new_draft)
        db.session.commit()
        
        # Clear session preview data
        session.pop('preview_data', None)
        clear_session_reserved_phones()
        
        return jsonify({
            'success': True,
            'message': 'Draft saved successfully',
            'draft_id': new_draft.id
        })
    
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/dashboard/stats')
@login_required
def api_dashboard_stats():
    """Get dashboard statistics for current user"""
    try:
        user_id = session.get('user_id')
        
        stats = {
            'drafts': Draft.query.filter_by(user_id=user_id, status='draft').count(),
            'pending': Draft.query.filter_by(user_id=user_id, status='pending').count(),
            'approved': Draft.query.filter_by(user_id=user_id, status='approved').count(),
            'generated': Draft.query.filter_by(user_id=user_id, status='generated').count()
        }
        
        return jsonify({'success': True, 'stats': stats})
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/save_draft', methods=['POST'])
@login_required
def save_draft():
    try:
        if 'preview_data' not in session:
            return jsonify({'success': False, 'message': 'No preview data found.'})
        
        preview_data = session['preview_data']
        user_id = session.get('user_id')
        
        template_config = TEMPLATE_CONFIG.get(preview_data['template_type'], {})
        
        new_draft = Draft(
            user_id=user_id,
            template_type=preview_data['template_type'],
            template_name=template_config.get('name', preview_data['template_type']),
            old_name=preview_data['replacements'].get('OLD_NAME', ''),
            replacements=preview_data['replacements'],
            preview_data=preview_data,
            status='draft'
        )
        
        db.session.add(new_draft)
        db.session.commit()
        
        session.pop('preview_data', None)
        clear_session_reserved_phones()
        
        return jsonify({
            'success': True,
            'message': 'Draft saved successfully!',
            'draft_id': new_draft.id
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

# Initialize database tables
with app.app_context():
    db.create_all()

# For local development
if __name__ == '__main__':
    app.run(debug=True)
